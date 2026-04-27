#!/usr/bin/env python3
"""
Build a styled DOCX from markdown for BUCT proposal delivery.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TABLE_ALIGN_RE = re.compile(r"^\s*:?-{3,}:?\s*$")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def set_run_font(run, name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_common(p, line_spacing=1.5, space_after=6) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)


def add_field(paragraph, field_code: str) -> None:
    r1 = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    r1._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    r1._r.append(instr_text)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    r1._r.append(fld_char_sep)

    paragraph.add_run(" ")

    r3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r3._r.append(fld_char_end)


def configure_document(doc: Document, header_text: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    header_p = section.header.paragraphs[0]
    header_p.text = header_text
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_p.runs:
        set_run_font(header_p.runs[0], "SimSun", 10)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer_p, "PAGE")
    if footer_p.runs:
        set_run_font(footer_p.runs[0], "SimSun", 10)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)

    for style_name, font_name, size in [
        ("Heading 1", "SimHei", 16),  # sanhao
        ("Heading 2", "SimHei", 15),  # xiaosanhao
        ("Heading 3", "SimHei", 14),  # sihao
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True


def add_inline_markdown_runs(paragraph, text: str, font: str, size: int) -> None:
    idx = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > idx:
            r = paragraph.add_run(text[idx : m.start()])
            set_run_font(r, font, size, bold=False)
        r = paragraph.add_run(m.group(1))
        set_run_font(r, font, size, bold=True)
        idx = m.end()
    if idx < len(text):
        r = paragraph.add_run(text[idx:])
        set_run_font(r, font, size, bold=False)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        raw = line.strip()
        if not raw:
            continue
        if not (raw.startswith("|") and raw.endswith("|")):
            continue
        cols = [c.strip() for c in raw.strip("|").split("|")]
        if cols and all(TABLE_ALIGN_RE.match(c) for c in cols):
            continue
        rows.append(cols)
    return rows


def write_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell_p = table.cell(r_idx, c_idx).paragraphs[0]
            cell_p.text = text
            set_paragraph_common(cell_p, line_spacing=1.2, space_after=0)
            if cell_p.runs:
                set_run_font(cell_p.runs[0], "SimSun", 11, bold=(r_idx == 0))


def consume_table(lines: list[str], start: int) -> tuple[int, list[str]]:
    out: list[str] = []
    i = start
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        if raw.strip().startswith("|") and raw.strip().endswith("|"):
            out.append(raw)
            i += 1
        else:
            break
    return i, out


def add_toc(doc: Document) -> None:
    p_title = doc.add_paragraph("Contents", style="Heading 1")
    set_paragraph_common(p_title)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    for run in p.runs:
        set_run_font(run, "SimSun", 12)
    doc.add_page_break()


def build_from_markdown(doc: Document, lines: Iterable[str]) -> None:
    src = list(lines)
    in_code = False
    i = 0
    while i < len(src):
        line = src[i].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            set_paragraph_common(p, line_spacing=1.1, space_after=0)
            r = p.add_run(line)
            set_run_font(r, "Consolas", 10)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            next_i, table_lines = consume_table(src, i)
            write_table(doc, parse_table_rows(table_lines))
            i = next_i
            continue

        h = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h:
            level = min(len(h.group(1)), 3)
            text = h.group(2).strip()
            p = doc.add_paragraph(text, style=f"Heading {level}")
            set_paragraph_common(p)
            if p.runs:
                size = 16 if level == 1 else 15 if level == 2 else 14
                set_run_font(p.runs[0], "SimHei", size, bold=True)
            i += 1
            continue

        if re.match(r"^\d+\.\s+.+$", stripped):
            p = doc.add_paragraph(style="List Number")
            set_paragraph_common(p)
            add_inline_markdown_runs(
                p,
                re.sub(r"^\d+\.\s+", "", stripped),
                "SimSun",
                12,
            )
            i += 1
            continue

        if re.match(r"^[-*]\s+.+$", stripped):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_common(p)
            add_inline_markdown_runs(
                p,
                re.sub(r"^[-*]\s+", "", stripped),
                "SimSun",
                12,
            )
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            set_paragraph_common(p)
            r = p.add_run(stripped.lstrip(">").strip())
            set_run_font(r, "KaiTi", 12)
            i += 1
            continue

        p = doc.add_paragraph()
        set_paragraph_common(p)
        p.paragraph_format.first_line_indent = Pt(24)
        add_inline_markdown_runs(p, stripped, "SimSun", 12)
        i += 1


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Input markdown path")
    parser.add_argument("--output", required=True, help="Output docx path")
    parser.add_argument("--profile", default="buct_default")
    parser.add_argument("--toc", default="true", choices=["true", "false"])
    parser.add_argument("--header", default="CoPilot Care (BUCT)")
    args = parser.parse_args()

    in_path = Path(args.input)
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    text = in_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc, args.header)

    if args.toc == "true":
        add_toc(doc)

    build_from_markdown(doc, text.splitlines())
    doc.save(str(out_path))
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
